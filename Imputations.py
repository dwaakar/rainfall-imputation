#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
make_all_figures.py
Generates 9 manuscript-ready, non-line/non-bar/non-scatter visualizations + a DOCX summary
for the rainfall imputation pipeline and results.

Inputs (put these in the working directory):
  - cv_results_metrics_all.csv
  - trend_preservation_all.csv
  - Station_deets.csv
  - daily rf1990-2023.csv
  - (optional) station_completeness_by_window.csv
  - (optional) imputed_monthly_totals.csv  # Station,Method,Month(YYYY-MM),Total

Usage:
  python make_all_figures.py --outdir outputs_all --year_for_matrix 2020 --window Full

Inspired by:
  - Galicia case (missingness matrices; CV flow)       [Vidal-Paz et al., 2023] 
  - Johor basin study (performance heat maps; PDFs)     [Sa’adi et al., 2023]
  - Susurluk basin (elevation stratification; homogeneity)[Hırca & Türkkan, 2024]
"""

import os
import sys
import argparse
import warnings
import numpy as np
import pandas as pd
import matplotlib
matplotlib.use("Agg")  # headless
import matplotlib.pyplot as plt
from matplotlib.patches import Rectangle, FancyArrow, FancyBboxPatch
from scipy.stats import gaussian_kde
from docx import Document
from docx.shared import Pt
warnings.filterwarnings("ignore")

# -----------------------------
# Utility / IO
# -----------------------------
def ensure_dir(path):
    os.makedirs(path, exist_ok=True)

def load_csv_safe(fname):
    if os.path.exists(fname):
        return pd.read_csv(fname)
    else:
        print(f"[WARN] Missing file: {fname}")
        return None

def assign_elev_band(elev):
    if pd.isna(elev): 
        return "unknown"
    if 0 <= elev < 800:      return "0-800m"
    if 800 <= elev < 1500:   return "800-1500m"
    if 1500 <= elev < 2500:  return "1500-2500m"
    if elev >= 2500:         return "2500-9999m"
    return "unknown"

def compute_completeness_by_window(rain_df, meta_df, windows):
    rows = []
    for st in [c for c in rain_df.columns if c != "Time"]:
        el = meta_df.loc[meta_df["Stations"].astype(str).str.strip()==str(st), "Elevation"]
        elev = el.iloc[0] if len(el)>0 else np.nan
        for lbl, (sy, ey) in windows.items():
            dwin = rain_df[(rain_df["Time"].dt.year>=sy) & (rain_df["Time"].dt.year<=ey)]
            comp = dwin[st].notna().mean() * 100.0
            rows.append({"Station": st, "Elevation": elev, "Band": assign_elev_band(elev), f"Completeness_{lbl}": comp})
    # merge rows by Station/Elevation
    out = pd.DataFrame(rows).groupby(["Station","Elevation","Band"]).agg("first").reset_index()
    return out

def pick_representatives(comp_df, window="Full", min_comp=90.0):
    # One representative station per band with completeness >= min_comp
    reps = {}
    for band, group in comp_df.groupby("Band"):
        cand = group[group[f"Completeness_{window}"] >= min_comp]
        if len(cand) == 0:
            continue
        # choose station with highest completeness, tie-break by elevation then station name
        cand = cand.sort_values([f"Completeness_{window}","Elevation","Station"], ascending=[False,False,True])
        reps[band] = cand.iloc[0]["Station"]
    return reps

# -----------------------------
# FIGURE 1: Missingness Pattern Matrix (tile)
# -----------------------------
def fig_missingness_matrix(rain_df, station, year, outpath):
    dyear = rain_df[rain_df["Time"].dt.year==year]
    if dyear.empty:
        print(f"[WARN] No data for year {year}, skipping missingness matrix for {station}.")
        return
    # Build month-day tile matrix: months rows, days columns
    months = sorted(dyear["Time"].dt.month.unique())
    max_days = 31
    mat = np.zeros((len(months), max_days), dtype=int)  # 1=obs, 0=missing
    for i,m in enumerate(months):
        dm = dyear[dyear["Time"].dt.month==m]
        # pad to max_days
        days = dm["Time"].dt.day.values
        vals = dm[station].values
        for d, v in zip(days, vals):
            mat[i, d-1] = 1 if pd.notna(v) else 0
    # plot as tiles (no axes ticks)
    fig, ax = plt.subplots(figsize=(10, 3))
    ax.imshow(mat, cmap=plt.cm.get_cmap("Greys"), aspect="auto", interpolation="nearest")
    ax.set_title(f"Missingness Pattern Matrix — {station} ({year})")
    ax.set_ylabel("Months")
    ax.set_xlabel("Days")
    ax.set_yticks([]); ax.set_xticks([])
    plt.tight_layout()
    fig.savefig(outpath, dpi=300)
    plt.close(fig)

# -----------------------------
# FIGURE 2: Method×Missingness×GapType Heatmaps (tiles)
# -----------------------------
def fig_method_missingness_heatmaps(cv_df, metric="NSE", outpath_prefix="heatmap"):
    # Tile heatmaps per gap type
    for g in ["random", "block"]:
        sub = cv_df[cv_df["GapType"]==g]
        if sub.empty: 
            print(f"[WARN] No CV data for gap type={g}, skipping heatmap.")
            continue
        # pivot: rows=Method, cols=MissingFrac, values=median metric
        pvt = sub.groupby(["Method","MissingFrac"])[metric].median().unstack("MissingFrac")
        # imshow tiles
        fig, ax = plt.subplots(figsize=(12, max(4, 0.5*len(pvt))))
        im = ax.imshow(pvt.values, cmap="viridis", aspect="auto", interpolation="nearest")
        ax.set_title(f"{metric} Heatmap — GapType={g}")
        ax.set_yticks(range(len(pvt.index))); ax.set_yticklabels(pvt.index)
        ax.set_xticks(range(len(pvt.columns))); ax.set_xticklabels([f"{int(c*100)}%" for c in pvt.columns])
        # colorbar (tile-based)
        cbar = fig.colorbar(im, ax=ax, shrink=0.8); cbar.set_label(metric)
        plt.tight_layout()
        fig.savefig(f"{outpath_prefix}_{g}.png", dpi=300)
        plt.close(fig)

# -----------------------------
# Treemap helper (squarify algorithm)
# -----------------------------
def squarify(values, x=0, y=0, width=1, height=1):
    # Minimal squarify: returns list of rectangles (x, y, w, h)
    # Ref: Bruls et al. (2000) squarified treemaps (basic implementation)
    rects = []
    if not values:
        return rects
    values = [float(v) for v in values if v>0]
    total = sum(values)
    values = [v/total*width*height for v in values]  # areas
    def layoutrow(row, w, x, y, h):
        area_sum = sum(row)
        if w >= h:  # horizontal slice
            row_height = area_sum / w
            cx = x
            for area in row:
                rw = area / row_height
                rects.append((cx, y, rw, row_height))
                cx += rw
            y += row_height; h -= row_height
        else:
            row_width = area_sum / h
            cy = y
            for area in row:
                rh = area / row_width
                rects.append((x, cy, row_width, rh))
                cy += rh
            x += row_width; w -= row_width
        return x,y,w,h
    def worst(row, w, h):
        if not row: return float("inf")
        s = sum(row)
        if w >= h:
            rheight = s/w
            return max([(max(row)/rheight)/(min(row)/rheight) if min(row)>0 else float("inf")])
        else:
            rwidth = s/h
            return max([(max(row)/rwidth)/(min(row)/rwidth) if min(row)>0 else float("inf")])
    row = []
    w = width; h = height
    for area in values:
        if not row:
            row.append(area)
        else:
            if worst(row+[area], w, h) <= worst(row, w, h):
                row.append(area)
            else:
                x,y,w,h = layoutrow(row, w, x, y, h)
                row = [area]
    if row:
        x,y,w,h = layoutrow(row, w, x, y, h)
    return rects

# -----------------------------
# FIGURE 3: Elevation-band Treemap (area blocks subdivided by winners)
# -----------------------------
def fig_elev_band_treemap(cv_df, comp_df, metric="NSE", outpath="treemap_elev_band.png"):
    # For each band, determine winning method (highest median metric)
    bands = []
    sizes = []
    labels = []
    # area sizing: total ValidPairs per band (proxy for evaluated volume)
    for band, grp in comp_df.groupby("Band"):
        stations_in_band = grp["Station"].unique().tolist()
        sub = cv_df[cv_df["Station"].isin(stations_in_band)]
        if sub.empty: 
            continue
        # choose winner by metric
        win = sub.groupby("Method")[metric].median().sort_values(ascending=False).index[0]
        valpairs = sub["ValidPairs"].sum()
        bands.append(band); sizes.append(max(valpairs, 1)); labels.append(f"{band}\n● {win}")
    # normalize and draw treemap
    rects = squarify(sizes, x=0, y=0, width=1, height=1)
    fig, ax = plt.subplots(figsize=(8, 6))
    for i,(x,y,w,h) in enumerate(rects):
        ax.add_patch(Rectangle((x,y), w,h, facecolor=plt.cm.tab20(i%20), edgecolor="white"))
        ax.text(x+w/2, y+h/2, labels[i], ha="center", va="center", fontsize=10, color="white")
    ax.set_title("Elevation-band Treemap — Winning Method per Band")
    ax.set_xticks([]); ax.set_yticks([])
    plt.tight_layout()
    fig.savefig(outpath, dpi=300)
    plt.close(fig)

# -----------------------------
# FIGURE 4: Waffle Charts for Trend Sign Agreement
# -----------------------------
def waffle_chart(value, n_rows=10, n_cols=10):
    total = n_rows*n_cols
    filled = int(round(value*total))
    grid = np.zeros((n_rows, n_cols))
    # fill row-major
    r,c = 0,0
    for k in range(filled):
        grid[r,c] = 1
        c += 1
        if c>=n_cols:
            c=0; r+=1
    return grid

def fig_trend_waffles(tr_df, outpath_prefix="waffle_trend"):
    # Use mean sign agreement per method
    agg = tr_df.groupby("Method")["SenSignAgree"].mean().sort_values(ascending=False)
    for i,(method,val) in enumerate(agg.items()):
        grid = waffle_chart(val if pd.notna(val) else 0.0, 10, 10)
        fig, ax = plt.subplots(figsize=(3,3))
        ax.imshow(grid, cmap=plt.cm.get_cmap("Greens"), interpolation="nearest")
        ax.set_title(f"Trend Sign Agreement — {method}\n({val*100:.1f}%)")
        ax.set_xticks([]); ax.set_yticks([])
        plt.tight_layout()
        fig.savefig(f"{outpath_prefix}_{method}.png", dpi=300)
        plt.close(fig)

# -----------------------------
# FIGURE 5: Distribution Glyphs (“Density Strips”) for Monthly Totals
# -----------------------------
def monthly_totals(df, station, window):
    times = df["Time"]
    sy, ey = window
    dwin = df[(times.dt.year>=sy) & (times.dt.year<=ey)]
    m = dwin.groupby(dwin["Time"].dt.to_period("M"))[station].sum().dropna()
    # return month index (as integer 0..N-1), values
    return np.arange(len(m)), m.values

def density_strip(values, x0, y0, width=1.0, scale=1.0, ax=None, color=(0.2,0.4,0.8,0.6)):
    # KDE along x in [0,1], draw as filled polygon
    if len(values)<5: 
        return
    kde = gaussian_kde(values)
    xs = np.linspace(0,1,200)
    ys = kde(xs)
    ys = ys / ys.max() * scale
    poly_x = np.concatenate([x0 + xs*width, x0 + xs[::-1]*width])
    poly_y = np.concatenate([y0 + ys, y0 - ys[::-1]])
    ax.fill(poly_x, poly_y, color=color, edgecolor=None)

def fig_density_strips_observed(rain_df, reps, windows, outpath_prefix="density_strip"):
    # Observed-only strips for monthly totals; imputed optional (not available by default)
    for band, st in reps.items():
        # choose window (Full)
        win = windows["Full"]
        idx, vals = monthly_totals(rain_df, st, win)
        fig, ax = plt.subplots(figsize=(10,2))
        # pack strips along x
        # Fit one strip per year by grouping months into years
        # Here: single strip representing the distribution across all months in window
        density_strip(vals, x0=0, y0=0.5, width=1.0, scale=0.4, ax=ax, color=(0.3,0.6,0.9,0.6))
        ax.set_title(f"Monthly Totals Density Strip (Observed) — {st} [{band}]")
        ax.set_xticks([]); ax.set_yticks([])
        ax.set_xlim(0,1); ax.set_ylim(0,1)
        plt.tight_layout()
        fig.savefig(f"{outpath_prefix}_{band}_{st}.png", dpi=300)
        plt.close(fig)

# -----------------------------
# FIGURE 6: Process Flow Diagram (tiles + arrows)
# -----------------------------
def fig_flow_diagram(outpath="flow_diagram.png"):
    fig, ax = plt.subplots(figsize=(10,6))
    ax.set_xlim(0,10); ax.set_ylim(0,10)
    def box(x,y,w,h,text):
        rect = FancyBboxPatch((x,y), w,h, boxstyle="round,pad=0.2", facecolor="#4C72B0", edgecolor="white", alpha=0.9)
        ax.add_patch(rect)
        ax.text(x+w/2, y+h/2, text, ha="center", va="center", color="white", fontsize=10)
    # Stages
    box(0.5, 7.5, 2.5, 1.2, "Data & Metadata\nQC & Screening")
    box(4.0, 7.5, 2.5, 1.2, "Neighbor Selection\n(Pearson |r|, K=5)")
    box(7.5, 7.5, 2.5, 1.2, "CV Masks\n(Random/Block)")
    box(2.0, 4.5, 3.0, 1.2, "Imputation Methods\nLR / MLR / IDW / ElevIDW\nNR / SA / OK / CokrigElev / XGB")
    box(6.5, 4.5, 2.5, 1.2, "Metrics\nRMSE, MAE, R², NSE\nPBIAS, Precision/Recall")
    box(4.0, 1.5, 2.5, 1.2, "Trend Preservation\nKendall τ, Sen’s slope\nSign agree, slope bias")
    # Arrows
    def arrow(x1,y1,x2,y2):
        ax.add_patch(FancyArrow(x1,y1, x2-x1, y2-y1, width=0.05, length_includes_head=True, head_width=0.3, color="#999999"))
    arrow(3.0, 8.1, 4.0, 8.1)
    arrow(6.5, 8.1, 7.5, 8.1)
    arrow(8.8, 7.5, 8.0, 5.7)
    arrow(3.25, 7.5, 3.25, 5.7)
    arrow(5.0, 5.7, 7.5, 5.7)
    arrow(6.8, 4.5, 5.25, 2.7)
    ax.set_xticks([]); ax.set_yticks([])
    ax.set_title("Cross-Validation Pipeline (Tile/Area-based Flow)")
    plt.tight_layout()
    fig.savefig(outpath, dpi=300)
    plt.close(fig)

# -----------------------------
# FIGURE 7: Tile Calendar (Monsoon Rain/No-Rain)
# -----------------------------
def fig_tile_calendar_rain(rain_df, station, year, months=[6,7,8,9], threshold=0.1, outpath="tile_calendar.png"):
    d = rain_df[(rain_df["Time"].dt.year==year) & (rain_df["Time"].dt.month.isin(months))]
    if d.empty:
        print(f"[WARN] No monsoon data for {station} in {year}. Skipping tile calendar.")
        return
    # build grid: rows=months, cols=days (max 31)
    mat = np.zeros((len(months), 31), dtype=int)
    for i,m in enumerate(months):
        dm = d[d["Time"].dt.month==m]
        for _, row in dm.iterrows():
            day = row["Time"].day - 1
            val = row[station]
            mat[i, day] = 1 if (pd.notna(val) and val>threshold) else 0
    fig, ax = plt.subplots(figsize=(10,3))
    ax.imshow(mat, cmap=plt.cm.get_cmap("Blues"), aspect="auto", interpolation="nearest")
    ax.set_title(f"Tile Calendar — Rain Days (Observed) — {station} ({year})")
    ax.set_xticks([]); ax.set_yticks([])
    plt.tight_layout()
    fig.savefig(outpath, dpi=300)
    plt.close(fig)

# -----------------------------
# FIGURE 8: Adjacency Matrix of Station Correlations
# -----------------------------
def fig_adj_corr_matrix(rain_df, stations, window, outpath="adj_corr_matrix.png"):
    sy, ey = window
    dwin = rain_df[(rain_df["Time"].dt.year>=sy) & (rain_df["Time"].dt.year<=ey)]
    stations = [s for s in stations if s in dwin.columns and s!="Time"]
    if len(stations)==0: 
        print("[WARN] No stations for adjacency correlation matrix.")
        return
    mat = np.zeros((len(stations), len(stations)))
    for i,si in enumerate(stations):
        for j,sj in enumerate(stations):
            s = dwin[[si,sj]].dropna()
            r = np.nan
            if len(s)>=30:
                r = s[si].corr(s[sj])
            mat[i,j] = np.abs(r) if pd.notna(r) else 0.0
    # sort by avg corr
    avg = mat.mean(axis=1)
    order = np.argsort(-avg)
    mat = mat[order][:,order]
    stations_ord = [stations[k] for k in order]
    fig, ax = plt.subplots(figsize=(8,8))
    im = ax.imshow(mat, cmap="magma", aspect="equal", interpolation="nearest")
    ax.set_title("Adjacency Matrix — |Pearson r| among Stations")
    ax.set_xticks([]); ax.set_yticks([])
    cbar = fig.colorbar(im, ax=ax, shrink=0.7); cbar.set_label("|r|")
    plt.tight_layout()
    fig.savefig(outpath, dpi=300)
    plt.close(fig)

# -----------------------------
# FIGURE 9: Elevation Raster-Style Grid Map (tiles)
# -----------------------------
def fig_elev_grid_map(meta_df, comp_df, reps, outpath="elev_grid_map.png"):
    # normalize lat/lon to [0,1]
    m = meta_df.copy()
    m["Stations"] = m["Stations"].astype(str).str.strip()
    lat = m["Latitude"].astype(float); lon = m["Longitude"].astype(float)
    latn = (lat - lat.min())/(lat.max()-lat.min()+1e-9)
    lonn = (lon - lon.min())/(lon.max()-lon.min()+1e-9)
    # draw grid tiles
    fig, ax = plt.subplots(figsize=(8,6))
    for i,row in m.iterrows():
        x = (lonn.iloc[i]); y = (latn.iloc[i])
        w = 0.02; h = 0.02
        st = str(row["Stations"])
        # find band from comp_df
        band = comp_df.loc[comp_df["Station"]==st, "Band"]
        band = band.iloc[0] if len(band)>0 else "unknown"
        color = {"0-800m":"#a6cee3", "800-1500m":"#1f78b4", "1500-2500m":"#33a02c", "2500-9999m":"#b2df8a", "unknown":"#aaaaaa"}.get(band, "#aaaaaa")
        ax.add_patch(Rectangle((x,y), w, h, facecolor=color, edgecolor="white"))
        if st in reps.values():
            ax.add_patch(Rectangle((x,y), w, h, fill=False, edgecolor="yellow", linewidth=2))
    ax.set_title("Elevation Bands & Representative Stations (Tile Map)")
    ax.set_xticks([]); ax.set_yticks([])
    plt.tight_layout()
    fig.savefig(outpath, dpi=300)
    plt.close(fig)

# -----------------------------
# DOCX summary (re-create)
# -----------------------------
def write_docx_summary(outpath="code_step_by_step_summary.docx"):
    doc = Document()
    title = doc.add_paragraph()
    run = title.add_run("Step-by-Step Summary of the Rainfall Imputation Code")
    run.bold = True; run.font.size = Pt(16)

    doc.add_paragraph(
        "This document summarizes the workflow implemented in Day10_rewrite.py, including data preparation, "
        "neighbor selection, simulated-gap cross-validation, imputation methods, scoring, and trend analysis."
    )

    doc.add_heading("1. Inputs & Outputs", level=1)
    doc.add_paragraph("• Inputs: daily rainfall with 'Time' + station columns; station metadata with station, lat, lon, elevation.")
    doc.add_paragraph("• Outputs: station_completeness_by_window.csv; cv_results_metrics_all.csv; trend_preservation_all.csv.")

    doc.add_heading("2. QC & Screening", level=1)
    doc.add_paragraph("1) Clamp negatives to 0; set extreme >500 mm to NaN.")
    doc.add_paragraph("2) Drop stations with 100% missing; compute completeness per window and assign elevation bands.")

    doc.add_heading("3. Neighbor Selection", level=1)
    doc.add_paragraph("Pearson correlation on overlapping days (≥60); select top-5 neighbors by |r|.")

    doc.add_heading("4. Simulated-Gap Cross-Validation", level=1)
    doc.add_paragraph("Random vs block (7/15/30), missing fractions 10–90%; strict test mask (score only where neighbors present at simulated gaps).")

    doc.add_heading("5. Imputation Methods", level=1)
    doc.add_paragraph("LR; MLR (anomaly + PCA when >1 neighbor); IDW; ElevIDW (exp(-|Δz|/z0)); Normal Ratio; Simple Average; Ordinary Kriging (Spherical variogram; fallback to IDW); CokrigElev (ElevIDW surrogate); XGBoost (neighbors, distance, elevation Δ, month).")

    doc.add_heading("6. Metrics & Trend", level=1)
    doc.add_paragraph("RMSE, MAE, R², NSE, PBIAS, Precision/Recall, ValidPairs; monthly Kendall τ; Sen’s slope; slope bias; sign agreement.")

    doc.add_heading("7. Elevation Use", level=1)
    doc.add_paragraph("Explicit in ElevIDW/CokrigElev weights and XGBoost features; not used in MLR, NR, SA, LR, current OK. Elevation bands for reporting only.")
    doc.save(outpath)
    print(f"[INFO] DOCX written: {outpath}")

# -----------------------------
# Main
# -----------------------------
def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--outdir", default="outputs_all", help="Output directory")
    parser.add_argument("--year_for_matrix", type=int, default=2020, help="Year for missingness matrix & tile calendar")
    parser.add_argument("--window", default="Full", choices=["Full","Last20","Last10","Last5"], help="Evaluation window for representatives")
    args = parser.parse_args()

    outdir = args.outdir
    figdir = os.path.join(outdir, "figs")
    docdir = os.path.join(outdir, "docx")
    ensure_dir(outdir); ensure_dir(figdir); ensure_dir(docdir)

    # Load inputs
    cv = load_csv_safe("cv_results_metrics_all.csv")
    tr = load_csv_safe("trend_preservation_all.csv")
    meta = load_csv_safe("Station_deets.csv")
    rain = load_csv_safe("daily rf1990-2023.csv")

    if rain is not None:
        rain["Time"] = pd.to_datetime(rain["Time"], errors="coerce")

    # Windows
    windows = {
        "Full": (1990, 2023),
        "Last20": (2003, 2023),
        "Last10": (2013, 2023),
        "Last5":  (2018, 2023),
    }

    # Completeness table
    comp = load_csv_safe("station_completeness_by_window.csv")
    if comp is None and (rain is not None and meta is not None):
        comp = compute_completeness_by_window(rain, meta, windows)
        comp.to_csv(os.path.join(outdir, "station_completeness_by_window.csv"), index=False)
        print(f"[INFO] Computed and saved station_completeness_by_window.csv in {outdir}")

    if comp is None:
        print("[WARN] No completeness info; some figures will be skipped.")
        comp = pd.DataFrame(columns=["Station","Elevation","Band"])

    # Representatives per elevation band
    reps = pick_representatives(comp, window=args.window, min_comp=90.0)
    if len(reps)==0:
        print("[WARN] No representative stations found at ≥90% completeness. Consider lowering threshold or recomputing completeness.")

    # FIGURE 1: Missingness matrices for representatives
    if rain is not None and len(reps)>0:
        for band, st in reps.items():
            fig_missingness_matrix(rain, st, args.year_for_matrix, os.path.join(figdir, f"missingness_matrix_{band}_{st}_{args.year_for_matrix}.png"))

    # FIGURE 2: Heatmaps (Method×Missingness×GapType)
    if cv is not None:
        # Ensure numeric
        for c in ["RMSE","NSE","ValidPairs","MissingFrac"]:
            if c in cv.columns: cv[c] = pd.to_numeric(cv[c], errors="coerce")
        fig_method_missingness_heatmaps(cv, metric="NSE", outpath_prefix=os.path.join(figdir, "heatmap_NSE"))
        fig_method_missingness_heatmaps(cv, metric="RMSE", outpath_prefix=os.path.join(figdir, "heatmap_RMSE"))

    # FIGURE 3: Elevation-band Treemap
    if cv is not None and not comp.empty:
        fig_elev_band_treemap(cv, comp, metric="NSE", outpath=os.path.join(figdir, "treemap_elev_band.png"))

    # FIGURE 4: Waffle charts (trend sign agreement)
    if tr is not None and "SenSignAgree" in tr.columns:
        tr["SenSignAgree"] = pd.to_numeric(tr["SenSignAgree"], errors="coerce")
        fig_trend_waffles(tr, outpath_prefix=os.path.join(figdir, "waffle_trend"))

    # FIGURE 5: Density strips (observed only; imputed optional)
    if rain is not None and len(reps)>0:
        fig_density_strips_observed(rain, reps, windows, outpath_prefix=os.path.join(figdir, "density_strip"))

    # FIGURE 6: Flow diagram
    fig_flow_diagram(outpath=os.path.join(figdir, "flow_diagram.png"))

    # FIGURE 7: Tile calendar (monsoon rain)
    if rain is not None and len(reps)>0:
        for band, st in reps.items():
            fig_tile_calendar_rain(rain, st, args.year_for_matrix, months=[6,7,8,9], outpath=os.path.join(figdir, f"tile_calendar_{band}_{st}_{args.year_for_matrix}.png"))

    # FIGURE 8: Adjacency correlation matrix (use up to 40 stations to keep readable)
    if rain is not None:
        stations = [c for c in rain.columns if c!="Time"][:40]
        fig_adj_corr_matrix(rain, stations, windows.get(args.window, windows["Full"]), outpath=os.path.join(figdir, "adj_corr_matrix.png"))

    # FIGURE 9: Elevation grid map
    if meta is not None and not comp.empty:
        fig_elev_grid_map(meta, comp, reps, outpath=os.path.join(figdir, "elev_grid_map.png"))

    # DOCX summary
    write_docx_summary(outpath=os.path.join(docdir, "code_step_by_step_summary.docx"))
    print(f"[DONE] All requested outputs saved in: {outdir}")

if __name__=="__main__":
    main()